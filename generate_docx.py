import docx
from docx.shared import Inches

doc = docx.Document()
doc.add_heading('Hướng dẫn sử dụng: Chức năng Danh Mục', 0)

doc.add_paragraph('Tài liệu này hướng dẫn cách truy cập và sử dụng các menu trong chức năng "Danh Mục" trên hệ thống PMS.')

doc.add_heading('1. Bán Thành Phẩm', level=1)
doc.add_paragraph('Bước 1: Mở thanh menu bên trái (Left Navigation) và mở rộng thư mục "Danh Mục".')
doc.add_paragraph('Bước 2: Click vào mục "Bán Thành Phẩm". Màn hình quản lý Bán thành phẩm sẽ hiển thị danh sách tất cả mã bán thành phẩm, quy cách, trạng thái và các nút chức năng để thêm mới, chỉnh sửa, hoặc xóa.')
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\intermediate_cat_1782986501375.png', width=Inches(6.0))

doc.add_heading('2. Thành Phẩm', level=1)
doc.add_paragraph('Bước 1: Trong menu "Danh Mục", click vào mục "Thành Phẩm".')
doc.add_paragraph('Bước 2: Màn hình hiển thị danh sách các Thành phẩm, cho phép bạn thực hiện các thao tác quản lý tương tự Bán thành phẩm nhưng áp dụng cho mã thành phẩm cuối cùng.')
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\product_cat_1782986536861.png', width=Inches(6.0))

doc.add_heading('3. BT - HC B1 / B2', level=1)
doc.add_paragraph('Đây là khu vực quản lý danh mục thiết bị và tiện ích cho từng phân xưởng riêng biệt (ví dụ: B1, B2). Bạn mở rộng mục "BT - HC B1" hoặc "BT - HC B2" để thấy các tùy chọn bên trong.')

doc.add_heading('3.1. Hiệu Chuẩn', level=2)
doc.add_paragraph('Click vào "Hiệu Chuẩn" để xem, thêm, sửa, xóa các thiết bị đo lường cần theo dõi hiệu chuẩn theo định kỳ tại phân xưởng tương ứng.')
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\hieuchuan_b1_1782986559513.png', width=Inches(6.0))

doc.add_heading('3.2. Bảo Trì', level=2)
doc.add_paragraph('Click vào "Bảo Trì" để quản lý các danh mục máy móc thiết bị chính thức cần được thực hiện bảo trì, bảo dưỡng của phân xưởng.')
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\baotri_b1_1782986581673.png', width=Inches(6.0))

doc.add_heading('3.3. Tiện Ích', level=2)
doc.add_paragraph('Click vào "Tiện Ích" để thao tác với danh sách các hệ thống phụ trợ (điện, nước, khí nén, HVAC,...) hỗ trợ cho phân xưởng.')
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\tienich_b1_1782986600020.png', width=Inches(6.0))

doc.save(r'C:\Users\QA2-Phongnh\Desktop\HDSD_DanhMuc_PMS.docx')
