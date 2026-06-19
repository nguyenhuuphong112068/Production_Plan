import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def main():
    doc = docx.Document()
    
    # Setup styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    run = title.add_run('HƯỚNG DẪN THAO TÁC KỸ THUẬT: CHỨC NĂNG CẢNH BÁO LỊCH')
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(20)

    img_path = r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\21a5e019-fe82-45a5-812a-21b94ea9bbf4\media__1781840831141.png'
    
    # 1. Phân xưởng
    h1 = doc.add_heading('1. THAO TÁC DÀNH CHO PHÂN XƯỞNG (ĐỀ NGHỊ)', level=1)
    
    doc.add_paragraph('Truy cập menu: Kế hoạch -> Cảnh báo lịch. Tại màn hình này, thực hiện tuần tự các bước:')
    
    p_step1 = doc.add_paragraph()
    p_step1.add_run('Bước 1: Chọn lô bị cảnh báo').bold = True
    doc.add_paragraph('• Mở Tab "Danh sách không đáp ứng KSC" hoặc "Cảnh báo ngày đáp ứng NL/BB".\n• Tích chọn (☑) vào ô checkbox đầu dòng của lô hàng cần xin ý kiến lùi ngày.', style='List Bullet')
    
    try:
        doc.add_picture(img_path, width=Inches(6.0))
        p_cap1 = doc.add_paragraph('Hình: Màn hình Cảnh báo lịch - Các Tab và nút Đề nghị')
        p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap1.runs[0].font.italic = True
        p_cap1.runs[0].font.size = Pt(9)
    except:
        pass

    p_step2 = doc.add_paragraph()
    p_step2.add_run('\nBước 2: Gửi Đề nghị').bold = True
    doc.add_paragraph('• Click nút xanh "Đề nghị chấp nhận ngày đáp ứng" (hoặc "Đề nghị chấp nhận thay đổi NL/BB") ở trên cùng bảng dữ liệu.\n• Lô hàng sẽ được gắn nhãn xanh "Đã đề nghị thay đổi ngày".', style='List Bullet')

    p_step3 = doc.add_paragraph()
    p_step3.add_run('\nBước 3: Gửi tin nhắn giải trình').bold = True
    doc.add_paragraph('• Tại cột "Trao đổi thông tin" của lô vừa đề xuất, nhập lý do (VD: Hỏng máy, thiếu vật tư...).\n• Click nút "Gửi".', style='List Bullet')

    doc.add_paragraph('\n' + '-'*60 + '\n')

    # 2. Phòng kế hoạch
    h2 = doc.add_heading('2. THAO TÁC DÀNH CHO PHÒNG KẾ HOẠCH (XEM XÉT)', level=1)
    
    p_step4 = doc.add_paragraph()
    p_step4.add_run('Bước 1: Tiếp nhận Đề nghị').bold = True
    doc.add_paragraph('• Truy cập menu: Kế hoạch -> Cảnh báo lịch.\n• Mở Tab "Đề Nghị Đổi Ngày KCS" hoặc "Xem Xét Đổi Ngày NL/BB".', style='List Bullet')

    p_step5 = doc.add_paragraph()
    p_step5.add_run('\nBước 2: Kiểm tra lý do').bold = True
    doc.add_paragraph('• Đọc nội dung giải trình của Phân xưởng tại cột "Trao đổi thông tin".\n• Nhập tin nhắn phản hồi (nếu cần) và bấm "Gửi".', style='List Bullet')

    p_step6 = doc.add_paragraph()
    p_step6.add_run('\nBước 3: Thực hiện Phê duyệt').bold = True
    doc.add_paragraph('Tại cột "Hành Động", click một trong hai nút xử lý:', style='Normal')
    
    doc.add_paragraph('• CHẤP NHẬN: Nhập ngày dự kiến mới vào ô trống. Lô hàng sẽ được cập nhật ngày mới, mất dòng báo lỗi đỏ, cho phép Phân xưởng Submit Lịch.', style='List Bullet')
    doc.add_paragraph('• KHÔNG CHẤP NHẬN: Nhập lý do từ chối vào cửa sổ hiển thị. Lô hàng vẫn giữ nguyên báo lỗi, Phân xưởng phải dời lịch lại.', style='List Bullet')

    out_path = r'c:\PMS\Production_Plan\HDSD_Canh_Bao_Lich_Ky_Thuat.docx'
    doc.save(out_path)
    print(f"Created: {out_path}")

if __name__ == '__main__':
    main()
