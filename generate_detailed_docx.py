import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('HƯỚNG DẪN SỬ DỤNG: CHỨC NĂNG DANH MỤC', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Tài liệu này hướng dẫn chi tiết cách sử dụng các tính năng trong hệ thống quản lý Danh Mục, giải thích chức năng của từng nút bấm và thao tác trên từng màn hình cụ thể.')
doc.add_paragraph('-' * 50)

# 1. Bán Thành Phẩm
doc.add_heading('1. DANH MỤC BÁN THÀNH PHẨM', level=1)
doc.add_paragraph('Màn hình này dùng để quản lý toàn bộ các mã Bán Thành Phẩm (BTP) hiện có trong hệ thống.')
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\intermediate_cat_1782986501375.png', width=Inches(6.0))

doc.add_heading('Chi tiết các nút chức năng:', level=2)
doc.add_paragraph('• Nút "Thêm Nguyên Liệu Mới" (Màu xanh lá):')
doc.add_paragraph('  - Tác dụng: Tạo mới một Bán thành phẩm.')
doc.add_paragraph('  - Cách thực hiện: Bấm vào nút, nhập các thông tin BTP vào cửa sổ mở ra. Sau đó lưu lại.')

doc.add_paragraph('• Nút "Sửa" (Màu vàng, chữ "Sửa"):')
doc.add_paragraph('  - Tác dụng: Chỉnh sửa thông tin của một BTP đã tồn tại.')
doc.add_paragraph('  - Cách thực hiện: Tìm dòng tương ứng với mã BTP cần sửa, bấm vào nút màu vàng. Sửa thông tin và lưu.')

doc.add_paragraph('• Nút "Công thức" (Màu xanh dương):')
doc.add_paragraph('  - Tác dụng: Thiết lập công thức BOM cho BTP.')
doc.add_paragraph('  - Cách thực hiện: Bấm vào nút để chuyển sang trang nhập liệu các thành phần cấu thành sản phẩm.')

doc.add_paragraph('• Nút trạng thái (Màu xanh dương có chữ Active/Deactive):')
doc.add_paragraph('  - Tác dụng: Bật hoặc tắt trạng thái sử dụng của mã BTP.')

doc.add_paragraph('• Nút "Lịch sử" (Biểu tượng đồng hồ):')
doc.add_paragraph('  - Tác dụng: Xem lịch sử cập nhật thay đổi của BTP.')


# 2. Thành Phẩm
doc.add_heading('2. DANH MỤC THÀNH PHẨM', level=1)
doc.add_paragraph('Quản lý thông tin Thành Phẩm cuối cùng của nhà máy.')
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\product_cat_1782986536861.png', width=Inches(6.0))

doc.add_heading('Chi tiết các nút chức năng:', level=2)
doc.add_paragraph('• Nút "Thêm Mới" / "Sửa": Cách dùng tương tự như Bán thành phẩm. Các trường thông tin bao gồm mã Thành Phẩm, quy cách hộp, liên kết với BTP nào, v.v.')
doc.add_paragraph('• Nút "Công thức Bao Bì": Thiết lập định mức màng, chai, vỉ cho riêng mã Thành Phẩm đó.')


# 3. Bảo trì - Hiệu chuẩn
doc.add_heading('3. QUẢN LÝ THIẾT BỊ / TIỆN ÍCH (BT - HC B1 & B2)', level=1)

doc.add_heading('3.1. Hiệu Chuẩn', level=2)
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\hieuchuan_b1_1782986559513.png', width=Inches(6.0))
doc.add_paragraph('• Nút "Thêm Mới": Thêm một dụng cụ đo lường cần hiệu chuẩn định kỳ. Người dùng nhập tên, tần suất và lịch dự kiến.')
doc.add_paragraph('• Nút "Sửa": Chỉnh sửa thông tin đo lường.')
doc.add_paragraph('• Nút "Xóa" (Màu đỏ): Bấm vào để loại bỏ dụng cụ đo (khi không còn sử dụng/hư hỏng).')

doc.add_heading('3.2. Bảo Trì', level=2)
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\baotri_b1_1782986581673.png', width=Inches(6.0))
doc.add_paragraph('• Nút "Thêm mới / Sửa / Xóa": Cấu hình và thêm các máy móc vào danh sách bảo trì của bộ phận kỹ thuật để tạo lịch dễ dàng.')

doc.add_heading('3.3. Tiện Ích', level=2)
doc.add_picture(r'C:\Users\QA2-Phongnh\.gemini\antigravity-ide\brain\6823b7dd-4284-42c2-a19c-ef7c338cb2b6\tienich_b1_1782986600020.png', width=Inches(6.0))
doc.add_paragraph('• Hệ thống hoạt động tương tự: Nút Thêm Mới, Sửa, Xóa để cập nhật danh sách thiết bị tiện ích HVAC, nước, khí nén, điện.')

doc.save(r'C:\Users\QA2-Phongnh\Desktop\HDSD_DanhMuc_PMS_ChiTiet.docx')
